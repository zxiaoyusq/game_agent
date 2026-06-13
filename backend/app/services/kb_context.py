"""Agent 通用：知识库 → LLM 上下文 + 多模态消息构造。

任何 Agent（策划 / 美术 / 研发 / ...）都可以复用：
1) resolve_kb(kb_refs)           —— 读取所选知识库条目 / 文件夹，得到统一的 text + images
2) build_messages(...)           —— 拼 system + human 消息；vision 注入按 provider 自动判断
3) extract_text(chunk)           —— 流式 chunk 拆包，兼容 str 与 list[part] 两种 content

引用 key 协议（前后端共用，前端 makeKbKey/parseKbKey 应保持一致）：
- 文件引用： "file:<module>:<category>:<item_id>"
            （兼容旧三段 "<module>:<category>:<item_id>"）
- 文件夹引用："folder:<module>:<category>:<folder_path>"，folder_path 用 / 分隔，
            "" 表示分类根目录（递归整个分类）

设计意图：
- Agent 业务层只关心提示词内容，不再关心怎么读文件 / 处理图片 / 展开文件夹
- provider 感知集中在本模块；以后接 GPT-4V / Gemini Vision 时只动这里
- 知识库读取规则（单文件 4k 截断、folder 递归 60 份封顶等）也集中在这里
"""

from __future__ import annotations

import base64
import logging
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any, Dict, List, Optional, Tuple

from langchain_core.messages import BaseMessage, HumanMessage, SystemMessage

from app.config import LLM_MODEL_INDEX
from app.models.knowledge import KbItem
from app.services.knowledge import (
    get_folder,
    get_item,
    get_raw_path,
    list_items_under_folder,
)

logger = logging.getLogger(__name__)


# ---- 常量 -------------------------------------------------------------------

# 单文档纯文本截断阈值；超过的部分丢弃，并在末尾追加截断提示
DOC_TEXT_LIMIT = 4000

# 视为文本可读的扩展名集合；其它扩展名走 content_type 兜底判断
TEXTLIKE_EXTS = {".md", ".txt", ".json", ".csv", ".js", ".ts", ".py", ".html", ".css"}

# Office Word 文档扩展名（仅 .docx；.doc 旧二进制走元信息占位）
DOCX_EXTS = {".docx"}

# 图片扩展名集合
IMAGE_EXTS = {".png", ".jpg", ".jpeg", ".gif", ".webp"}

# 当前支持 vision 的 provider 集合；以后扩展时往这里加即可
VISION_PROVIDERS = {"anthropic"}

# folder 引用一次最多展开的文件总数（across 整个 resolve_kb 调用，全部 folder ref 共享一个池）
# 超出时尾部追加 "…还有 M 份未展开" 提示，避免悄悄丢资料
FOLDER_EXPAND_LIMIT = 60


# ---- 数据结构 ---------------------------------------------------------------

@dataclass
class KbContext:
    """resolve_kb 的返回值，封装文本块 + 图片 base64 列表。

    text:    给所有模型看的纯文本块（图片也会写一条文本占位，方便无 vision 模型 fallback）
    images:  仅供 vision 模型注入的 base64 图片，元素结构 {name, media_type, data}
    """

    text: str = ""
    images: List[Dict[str, Any]] = field(default_factory=list)


# ---- key 解析 ---------------------------------------------------------------

def _parse_ref(key: str) -> Optional[Tuple[str, str, str, str]]:
    """把引用 key 解析为 (kind, module, category, locator)，非法返回 None。

    locator：
    - file ref：item_id
    - folder ref：folder_path（"" 表示根目录）
    """
    if not key or not isinstance(key, str):
        return None
    parts = key.split(":")
    # 新协议：file:module:category:item_id
    if len(parts) >= 4 and parts[0] == "file":
        return ("file", parts[1], parts[2], ":".join(parts[3:]))
    # 新协议：folder:module:category:[folder_path...]
    if len(parts) >= 3 and parts[0] == "folder":
        # folder_path 内本身就不允许含冒号，但还是用 join 保底
        path = ":".join(parts[3:]) if len(parts) > 3 else ""
        return ("folder", parts[1], parts[2], path)
    # 兼容老协议：module:category:item_id
    if len(parts) == 3:
        return ("file", parts[0], parts[1], parts[2])
    return None


# ---- 单文件展开（file ref + folder ref 内部都用这个）----------------------

def _expand_item(
    item: KbItem,
    *,
    text_parts: List[str],
    images: List[Dict[str, Any]],
    title_prefix: str = "",
) -> None:
    """把单个 KbItem 展开为文本片段 + 可选的 base64 图片。

    title_prefix 仅用于在 head 行加前缀（例如 folder ref 展开时显示相对路径）。
    """
    module, category = item.module, item.category
    path = get_raw_path(module, category, item.id)

    ct = (item.content_type or "").lower()
    ext = Path(item.filename).suffix.lower()

    head_loc = (
        f"{module}/{category}/{item.folder} · {item.filename}"
        if item.folder
        else f"{module}/{category} · {item.filename}"
    )

    # 1) 文本类
    is_textlike = ext in TEXTLIKE_EXTS or ct.startswith("text/") or ct == "application/json"
    if is_textlike and path is not None:
        try:
            txt = path.read_text(encoding="utf-8", errors="ignore")
        except OSError:
            txt = ""
        if txt:
            if len(txt) > DOC_TEXT_LIMIT:
                txt = txt[:DOC_TEXT_LIMIT] + f"\n…（截断，原文 {len(txt)} 字）"
            head = f"{title_prefix}# {item.title}（{head_loc}）"
            if item.summary:
                head += f"\n摘要：{item.summary}"
            text_parts.append(f"{head}\n\n{txt}")
            return

    # 1.5) Word .docx：用 python-docx 抽段落 + 表格文本，再走通用截断
    if ext in DOCX_EXTS and path is not None:
        txt = _extract_docx_text(path)
        if txt:
            if len(txt) > DOC_TEXT_LIMIT:
                txt = txt[:DOC_TEXT_LIMIT] + f"\n…（截断，原文 {len(txt)} 字）"
            head = f"{title_prefix}# {item.title}（{head_loc}）"
            if item.summary:
                head += f"\n摘要：{item.summary}"
            text_parts.append(f"{head}\n\n{txt}")
            return
        # 抽不出文本（损坏 / 空白文档 / 解析失败）：fallback 到下面的元信息占位

    # 2) 图片类
    is_image = ct.startswith("image/") or ext in IMAGE_EXTS
    if is_image and path is not None:
        try:
            raw = path.read_bytes()
            media_type = ct if ct.startswith("image/") else _guess_image_media_type(ext)
            images.append(
                {
                    "name": item.filename,
                    "media_type": media_type,
                    "data": base64.b64encode(raw).decode("ascii"),
                }
            )
        except OSError:
            pass
        text_parts.append(
            f"{title_prefix}# 图片参考：{item.title}（{head_loc}）"
            + (f"\n摘要：{item.summary}" if item.summary else "")
        )
        return

    # 3) 其它多模态
    text_parts.append(
        f"{title_prefix}# 多模态参考：{item.title}（{head_loc}）"
        + (f"\n摘要：{item.summary}" if item.summary else "")
        + "\n（当前阶段未直接读取此类文件内容，只把元信息提供给模型作背景）"
    )


# ---- 知识库 → 上下文 --------------------------------------------------------

def resolve_kb(refs: Optional[List[str]]) -> KbContext:
    """根据引用 key 列表还原知识库内容。

    单个 key 可以是 file: 或 folder: 形式；folder 引用会递归展开下属所有文件，
    所有 folder 共享一个总配额 FOLDER_EXPAND_LIMIT，超出尾部追加未展开数量提示。

    多个引用条目之间用 \\n\\n---\\n\\n 分隔。
    """
    text_parts: List[str] = []
    images: List[Dict[str, Any]] = []

    # 跨 folder 共享的展开配额：避免选了多个大目录把上下文撑爆
    folder_expand_remaining = FOLDER_EXPAND_LIMIT

    for key in refs or []:
        parsed = _parse_ref(key)
        if parsed is None:
            # 非法 key 静默跳过，不影响其它引用
            continue
        kind, module, category, locator = parsed

        if kind == "file":
            item = get_item(module, category, locator)
            if item is None:
                continue
            _expand_item(item, text_parts=text_parts, images=images)
            continue

        # kind == "folder"
        try:
            sub_items = list_items_under_folder(module, category, locator)
        except ValueError:
            continue

        folder_meta = get_folder(module, category, locator) if locator else None

        # 文件夹自身的"头"段：无论是否能展开都给到提示词，让模型知道目录结构
        head_lines: List[str] = []
        loc_label = locator if locator else "(根目录)"
        head_lines.append(f"# 文件夹引用：{module}/{category}/{loc_label}")
        if folder_meta and folder_meta.desc:
            head_lines.append(f"描述：{folder_meta.desc}")
        head_lines.append(f"包含文件数：{len(sub_items)}")

        # 文件清单（无论是否展开内容都先把列表给到模型，至少知道有什么）
        if sub_items:
            file_lines = []
            for it in sub_items:
                rel = f"{it.folder}/{it.filename}" if it.folder else it.filename
                summary_part = f" — {it.summary}" if it.summary else ""
                file_lines.append(f"- {rel}（{it.title}{summary_part}）")
            head_lines.append("文件清单：\n" + "\n".join(file_lines))

        text_parts.append("\n".join(head_lines))

        # 按配额逐个展开
        if folder_expand_remaining <= 0:
            text_parts.append(
                f"（本次展开总配额已用尽，{len(sub_items)} 份文件均仅在清单中列出）"
            )
            continue

        to_expand = sub_items[:folder_expand_remaining]
        skipped = len(sub_items) - len(to_expand)
        for it in to_expand:
            _expand_item(
                it,
                text_parts=text_parts,
                images=images,
                title_prefix=f"[来自 {module}/{category}/{loc_label}] ",
            )
            folder_expand_remaining -= 1
            if folder_expand_remaining <= 0:
                break

        if skipped > 0:
            text_parts.append(
                f"（…还有 {skipped} 份未展开，受总配额 {FOLDER_EXPAND_LIMIT} 份限制）"
            )

    return KbContext(
        text="\n\n---\n\n".join(text_parts).strip(),
        images=images,
    )


def _extract_docx_text(path: Path) -> str:
    """从 .docx 抽段落 + 表格文本，返回纯文本（保留段间换行）。

    抽不出 / 解析失败时返回空串，让上层 fallback 到元信息占位。
    """
    try:
        # 延迟导入：python-docx 不算热路径，避免影响其它模块的启动
        from docx import Document  # type: ignore
    except ImportError:
        logger.warning("python-docx 未安装，跳过 docx 抽取：%s", path)
        return ""

    try:
        doc = Document(str(path))
    except Exception as e:
        logger.warning("docx 打开失败 %s：%s", path, e)
        return ""

    parts: List[str] = []

    # 段落：跳过空段
    for para in doc.paragraphs:
        text = (para.text or "").strip()
        if text:
            parts.append(text)

    # 表格：每行用 " | " 拼接成一行，行间换行；空表跳过
    for table in doc.tables:
        for row in table.rows:
            cells = [(c.text or "").strip().replace("\n", " ") for c in row.cells]
            row_text = " | ".join(c for c in cells if c)
            if row_text:
                parts.append(row_text)

    return "\n".join(parts).strip()


def _guess_image_media_type(ext: str) -> str:
    """按扩展名兜底，langchain anthropic 的 image part 必须有合法 media_type。"""
    return {
        ".png": "image/png",
        ".jpg": "image/jpeg",
        ".jpeg": "image/jpeg",
        ".gif": "image/gif",
        ".webp": "image/webp",
    }.get(ext, "image/png")


# ---- 消息构造（含 vision 注入） ---------------------------------------------

def build_messages(
    *,
    system: str,
    user_text: str,
    kb: Optional[KbContext] = None,
    model_id: str,
) -> List[BaseMessage]:
    """构造 [SystemMessage, HumanMessage]；按 provider 自动决定是否注入图片。

    - 若 kb.images 非空且 provider 在 VISION_PROVIDERS 中：
        HumanMessage.content = [{"type":"text", "text": user_text},
                                 {"type":"image", "source":{...}}, ...]
    - 否则：HumanMessage.content = user_text（纯字符串）

    user_text 由调用方自行拼装；本函数不假设其结构，只负责往里装多模态壳子。
    """
    images = (kb.images if kb else []) or []
    provider = LLM_MODEL_INDEX.get(model_id, {}).get("provider", "")
    use_vision = bool(images) and provider in VISION_PROVIDERS

    if use_vision:
        # langchain 的 HumanMessage content 支持 list[dict] 多模态格式
        # anthropic 接受 source.type=base64 + media_type
        parts: List[Dict[str, Any]] = [{"type": "text", "text": user_text}]
        for img in images:
            parts.append(
                {
                    "type": "image",
                    "source": {
                        "type": "base64",
                        "media_type": img["media_type"],
                        "data": img["data"],
                    },
                }
            )
        return [SystemMessage(content=system), HumanMessage(content=parts)]

    return [SystemMessage(content=system), HumanMessage(content=user_text)]


# ---- 流式拆包 ---------------------------------------------------------------

def extract_text(chunk: Any) -> Optional[str]:
    """从 AIMessageChunk / AIMessage 中抽纯文本，过滤工具调用等非文本片段。

    - content 为 str：直接返回
    - content 为 list：拼接所有 type=='text' 的 part 与裸字符串
    - 其它：返回 None
    """
    content = getattr(chunk, "content", None)
    if isinstance(content, str):
        return content or None
    if isinstance(content, list):
        parts: List[str] = []
        for item in content:
            if isinstance(item, dict) and item.get("type") == "text":
                parts.append(item.get("text", ""))
            elif isinstance(item, str):
                parts.append(item)
        return "".join(parts) or None
    return None
